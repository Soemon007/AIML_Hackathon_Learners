from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

def add_hyperlink(paragraph, url, text):

    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    
    rPr = OxmlElement('w:rPr')
    

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1') 
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')
    rPr.append(sz)
    
    new_run.append(rPr)
    
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink


def create_citations(private_data, public_data):
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.name = 'Arial'
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True
    heading1_style.font.color.rgb = RGBColor(0, 0, 0)
    
    heading2_style = doc.styles['Heading 2']
    heading2_style.font.name = 'Arial'
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True
    heading2_style.font.color.rgb = RGBColor(0, 0, 0)
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    title = doc.add_heading("Citations & Source Document", level=1)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    
    heading = doc.add_heading("Public Sources", level=2)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    
    source_urls = public_data.get("source_urls", [])
    
    if source_urls:
        intro = doc.add_paragraph(f"{len(source_urls)} public web pages reviewed:")
        intro.paragraph_format.space_after = Pt(6)

        for idx, url in enumerate(source_urls, start=1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            
            run = p.add_run(f"{idx}. ")
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            
            add_hyperlink(p, url, url)
    else:
        doc.add_paragraph("No public web sources used")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"output/citations/company_citations_{timestamp}.docx"
    
    doc.save(output_filename)
    print(f"💾 Saved citations to: {output_filename}")

